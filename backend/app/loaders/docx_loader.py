from pathlib import Path
from uuid import uuid4

from docx import Document as DocxDocument

from app.models.document import Document, Page
from app.utils.metadata import MetadataExtractor


class DOCXLoader:

    def load(self, docx_path: Path) -> Document:

        doc = DocxDocument(docx_path)

        text = "\n".join(
            paragraph.text
            for paragraph in doc.paragraphs
            if paragraph.text.strip()
        )

        metadata = MetadataExtractor().extract(
            docx_path.name,
            text,
        )

        return Document(
            document_id=str(uuid4()),
            filename=docx_path.name,
            total_pages=1,
            domain=metadata["domain"],
            department=metadata["department"],
            document_type="DOCX",
            metadata=metadata,
            pages=[
                Page(
                    page_number=1,
                    text=text,
                )
            ],
        )